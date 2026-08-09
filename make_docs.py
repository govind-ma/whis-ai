import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def convert_md(md_file, docx_file, pdf_file, title_text):
    # 1. WORD DOCX
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title_text)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(18, 52, 86)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Whis AI — Cyber Security & Scam Defense Platform | Android 7.0 to 16")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_str = line.strip()
        if not line_str: continue
        
        if line_str.startswith("# "):
            p = doc.add_heading(line_str[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(18, 52, 86)
        elif line_str.startswith("## "):
            p = doc.add_heading(line_str[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(30, 80, 130)
        elif line_str.startswith("### "):
            p = doc.add_heading(line_str[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(50, 100, 150)
        elif line_str.startswith("- ") or line_str.startswith("* "):
            doc.add_paragraph(line_str[2:], style='List Bullet')
        elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. ") or line_str.startswith("4. "):
            doc.add_paragraph(line_str, style='List Number')
        elif not line_str.startswith("```") and not line_str.startswith("|") and not line_str.startswith("---"):
            doc.add_paragraph(line_str)

    doc.save(docx_file)
    print(f"{docx_file} generated successfully!")

    # 2. REPORTLAB PDF
    pdf_doc = SimpleDocTemplate(pdf_file, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('PDFTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=16, leading=20, textColor=colors.HexColor('#123456'), alignment=1, spaceAfter=10)
    h1_style = ParagraphStyle('PDFH1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=13, leading=17, textColor=colors.HexColor('#1E5082'), spaceBefore=12, spaceAfter=6)
    h2_style = ParagraphStyle('PDFH2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#326496'), spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('PDFBody', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12.5, textColor=colors.HexColor('#222222'), spaceAfter=4)

    story = []
    story.append(Paragraph(title_text, title_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#123456'), spaceAfter=12))

    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("```") or line_str.startswith("|") or line_str.startswith("---"):
            continue
        
        clean_line = line_str.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        clean_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_line)
        clean_line = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', clean_line)
            
        try:
            if line_str.startswith("# "):
                story.append(Paragraph(clean_line[2:], h1_style))
            elif line_str.startswith("## "):
                story.append(Paragraph(clean_line[3:], h1_style))
            elif line_str.startswith("### "):
                story.append(Paragraph(clean_line[4:], h2_style))
            else:
                story.append(Paragraph(clean_line, body_style))
        except Exception:
            plain = line_str.replace("*", "").replace("`", "")
            story.append(Paragraph(plain, body_style))

    pdf_doc.build(story)
    print(f"{pdf_file} generated successfully!")

convert_md("d:\\1A\\WHIS_AI_PROJECT_REPORT.md", "d:\\1A\\WHIS_AI_PROJECT_REPORT.docx", "d:\\1A\\WHIS_AI_PROJECT_REPORT.pdf", "WHIS AI — COMPREHENSIVE CYBER DEFENSE PROJECT REPORT")
convert_md("d:\\1A\\CORE.md", "d:\\1A\\CORE.docx", "d:\\1A\\CORE.pdf", "WHIS AI — CORE MISSION, ARCHITECTURE & STRATEGY")
